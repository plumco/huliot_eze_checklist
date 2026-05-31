import os
import re
import io
import streamlit as st
from docx import Document

# Set page layout and design theme
st.set_page_config(
    page_title="Huliot India - Word Checklist Workstation",
    page_icon="💼",
    layout="wide"
)

# Custom branding CSS with correct Streamlit parameter
st.markdown("""
    <style>
    .brand-title {
        color: #2E7D32;
        font-family: 'Arial Black', sans-serif;
        font-size: 28px;
        font-weight: 900;
        letter-spacing: -1px;
        margin-bottom: 0px;
    }
    .brand-subtitle {
        color: #888888;
        font-size: 11px;
        font-weight: bold;
        letter-spacing: 3px;
        margin-top: 0px;
        margin-bottom: 20px;
    }
    .highlight-box {
        background-color: #FFFFF0;
        border: 1px solid #FFEB3B;
        padding: 12px;
        border-radius: 6px;
        margin-bottom: 15px;
    }
    </style>
""", unsafe_allow_html=True)

INIT_VALUES = {
    "date": "",
    "projectNo": "", 
    "salesPerson": "", 
    "salesPhone": "",
    "projectName": "", 
    "projectAddress": "", 
    "city": "",
    "clientName": "", 
    "clientPhone": "",
    "architectName": "", 
    "architectPhone": "",
    "mepName": "", 
    "mepPhone": "",

    # Drawing Checklist
    "drwTypical": False, 
    "drwToiletKitchen": False, 
    "drwPodium": False,
    "drwBuilding": False, 
    "drwArchOnly": False, 
    "drwOther": False, 
    "drwOtherText": "",

    # Drainage checklist variables
    "drnIntUS": False, 
    "drnIntHT": False,
    "drnVertUS": False, 
    "drnVertHT": False,
    "drnPodUS": False, 
    "drnPodHT": False,

    # Water supply checklist variables
    "wsIntPipePERT": False, 
    "wsIntPipePPR": False,
    "wsIntFitPPSU": False, 
    "wsIntFitBrass": False,
    "wsExtPipePERT": False, 
    "wsExtPipePPR": False,
    "wsExtFitPPSU": False, 
    "wsExtFitBrass": False,
    "wsTerPipePERT": False, 
    "wsTerPipePPR": False,
    "wsTerFitPPSU": False, 
    "wsTerFitBrass": False,

    # Building specs
    "priority": "", 
    "buildingType": "",
    "floors": "", 
    "shafts": "", 
    "rooms": "", 
    "floorHeight": "",
    "offerDate": "", 
    "probability": "50",

    # Drainage calculations style
    "calcDrainage": "", 
    "boqCompare": [], 
    "calcFor": "",

    # Routing checklist variables
    "rteCeiling": False,
    "rteToiletSunken": False, 
    "rteToiletMM": "",
    "rteKitchenSunken": False, 
    "rteKitchenMM": "",
    "rteUtilitySunken": False, 
    "rteUtilityMM": "",
    "wsCeilingDrop": False, 
    "wsWallChase": False,
    "serviceFloor": "", 
    "serviceFloorWhere": "",
    "notes": "",
}

# Initialize state structures
for key, val in INIT_VALUES.items():
    if key not in st.session_state:
        st.session_state[key] = val

st.markdown('<div class="brand-title">Huliot<span style="color:#555555; font-weight:300;">India</span></div>', unsafe_allow_html=True)
st.markdown('<div class="brand-subtitle">WE MAKE IT FLOW</div>', unsafe_allow_html=True)

st.title("Checklist Workstation & DOCX Exporter")
st.write("Fill in the fields below. The application will find your template and replace placeholders and toggle checkboxes verbatim.")

st.sidebar.header("System Controls")

# Create a clean side-by-side or full-width layout for the Reset action
if st.sidebar.button("🔄 Reset & Refresh Form", use_container_width=True, help="Wipes all entered inputs and reloads the screen clean."):
    for key in INIT_VALUES.keys():
        st.session_state[key] = INIT_VALUES[key]
    st.success("All fields cleared successfully!")
    st.rerun()

# Find the template file inside local workspace dynamically
local_templates = [f for f in os.listdir(".") if f.startswith("Checklist") and f.endswith(".docx")]
template_file_path = None

if local_templates:
    template_file_path = local_templates[0]
    st.sidebar.success(f"Detected committed file: `{template_file_path}`")
else:
    st.sidebar.warning("No committed Word template starting with 'Checklist' detected in root directory.")

# File uploader fallback option
uploaded_template = st.sidebar.file_uploader(
    "Optionally upload a different template file (.docx)", 
    type=["docx"], 
    help="Upload the official Huliot checklist template to preserve original fonts."
)

st.header("General Information")

col1, col2 = st.columns(2)
with col1:
    st.session_state.date = st.text_input("Date", value=st.session_state.date, placeholder="DD.MM.YYYY")
    st.session_state.projectNo = st.text_input("Project No. / Bitrix deal ID", value=st.session_state.projectNo)
    st.session_state.salesPerson = st.text_input("Huliot Sales Person Name", value=st.session_state.salesPerson)
    st.session_state.salesPhone = st.text_input("Sales Person Phone / No", value=st.session_state.salesPhone)
    st.session_state.projectName = st.text_input("Project Name & Address", value=st.session_state.projectName)

with col2:
    st.session_state.city = st.text_input("City", value=st.session_state.city)
    st.session_state.clientName = st.text_input("Customer / Client Name & No", value=st.session_state.clientName)
    st.session_state.architectName = st.text_input("Architect Name & No", value=st.session_state.architectName)
    st.session_state.mepName = st.text_input("MEP Consultant Name & No", value=st.session_state.mepName)

st.markdown("""
<div class="highlight-box">
    <strong>Material Composition: Consider</strong><br/>
    Kindly click on checkboxes below — Checked values will be represented as checked (☑) in your downloaded Word file without changing any alignments or formatting guidelines.
</div>
""", unsafe_allow_html=True)

st.subheader("Drawing attached in email:-")
col_drw1, col_drw2 = st.columns(2)
with col_drw1:
    st.session_state.drwTypical = st.checkbox("All or Typical Floor drawing with consultant/ client plumbing layout attached", value=st.session_state.drwTypical)
    st.session_state.drwToiletKitchen = st.checkbox("Typical Toilet and kitchen section details attached", value=st.session_state.drwToiletKitchen)
    st.session_state.drwPodium = st.checkbox("Podium/ Diversion/ Basement drawing attached", value=st.session_state.drwPodium)
with col_drw2:
    st.session_state.drwBuilding = st.checkbox("Full Building section view attached for Building height / podium/ till basement/Ground", value=st.session_state.drwBuilding)
    st.session_state.drwArchOnly = st.checkbox("Only Architectural layout attached", value=st.session_state.drwArchOnly)
    st.session_state.drwOther = st.checkbox("Any Other drawing attached", value=st.session_state.drwOther)

if st.session_state.drwOther:
    st.session_state.drwOtherText = st.text_input("Specify other drawing details", value=st.session_state.drwOtherText)

st.subheader("Drainage:")
col_drn1, col_drn2 = st.columns(2)
with col_drn1:
    st.session_state.drnIntUS = st.checkbox("Internal Toilets -Ultra Silent consider", value=st.session_state.drnIntUS)
    st.session_state.drnIntHT = st.checkbox("Internal Toilets -HT Pro consider", value=st.session_state.drnIntHT)
    st.session_state.drnVertUS = st.checkbox("Vertical/ External stack -Ultra Silent consider", value=st.session_state.drnVertUS)
with col_drn2:
    st.session_state.drnVertHT = st.checkbox("Vertical/ External stack -HT Pro consider", value=st.session_state.drnVertHT)
    st.session_state.drnPodUS = st.checkbox("Podium/Diversion/ Basement – Ultra Silent consider", value=st.session_state.drnPodUS)
    st.session_state.drnPodHT = st.checkbox("Podium/Diversion/ Basement - HT Pro consider", value=st.session_state.drnPodHT)

st.subheader("Water Supply:")
col_ws1, col_ws2, col_ws3 = st.columns(3)
with col_ws1:
    st.write("**Internal Pipe & Fittings**")
    st.session_state.wsIntPipePERT = st.checkbox("Internal Pipe (PERT/AL/PERT) Consider", value=st.session_state.wsIntPipePERT)
    st.session_state.wsIntPipePPR = st.checkbox("Internal Pipe (PPR) Consider", value=st.session_state.wsIntPipePPR)
    st.session_state.wsIntFitPPSU = st.checkbox("Internal Fittings PPSU Consider", value=st.session_state.wsIntFitPPSU)
    st.session_state.wsIntFitBrass = st.checkbox("Internal Fittings BRASS Consider", value=st.session_state.wsIntFitBrass)

with col_ws2:
    st.write("**External Pipe & Fittings**")
    st.session_state.wsExtPipePERT = st.checkbox("External Pipe (PERT/AL/PERT) Consider", value=st.session_state.wsExtPipePERT)
    st.session_state.wsExtPipePPR = st.checkbox("External Pipe (PPR) Consider", value=st.session_state.wsExtPipePPR)
    st.session_state.wsExtFitPPSU = st.checkbox("External Fittings PPSU Consider", value=st.session_state.wsExtFitPPSU)
    st.session_state.wsExtFitBrass = st.checkbox("External Fittings BRASS Consider", value=st.session_state.wsExtFitBrass)

with col_ws3:
    st.write("**Terrace looping Pipe & Fittings**")
    st.session_state.wsTerPipePERT = st.checkbox("Terrace looping Pipe (PERT/AL/PERT) Consider", value=st.session_state.wsTerPipePERT)
    st.session_state.wsTerPipePPR = st.checkbox("Terrace looping Pipe (PPR) Consider", value=st.session_state.wsTerPipePPR)
    st.session_state.wsTerFitPPSU = st.checkbox("Terrace looping Fittings PPSU Consider", value=st.session_state.wsTerFitPPSU)
    st.session_state.wsTerFitBrass = st.checkbox("Terrace looping Fittings BRASS Consider", value=st.session_state.wsTerFitBrass)

st.subheader("Priority & Building Information")
col_bld1, col_bld2 = st.columns(2)

with col_bld1:
    st.session_state.priority = st.selectbox("Priority", ["", "A: High", "B: Medium", "C: Low"], index=["", "A: High", "B: Medium", "C: Low"].index(st.session_state.priority) if st.session_state.priority in ["", "A: High", "B: Medium", "C: Low"] else 0)
    st.session_state.buildingType = st.selectbox("Kind of Building", [
        "", "Office / commercial building", "Shopping centre", "Hotel / Restaurant",
        "Industrial building", "Hospital / nursing home", "School building",
        "Social housing", "Sport facility", "Infrastructure buildings", "Residential building"
    ], index=[
        "", "Office / commercial building", "Shopping centre", "Hotel / Restaurant",
        "Industrial building", "Hospital / nursing home", "School building",
        "Social housing", "Sport facility", "Infrastructure buildings", "Residential building"
    ].index(st.session_state.buildingType) if st.session_state.buildingType in [
        "", "Office / commercial building", "Shopping centre", "Hotel / Restaurant",
        "Industrial building", "Hospital / nursing home", "School building",
        "Social housing", "Sport facility", "Infrastructure buildings", "Residential building"
    ] else 0)

with col_bld2:
    st.session_state.floors = st.text_input("No of Typical floors", value=st.session_state.floors)
    st.session_state.shafts = st.text_input("No of shafts", value=st.session_state.shafts)
    st.session_state.rooms = st.text_input("No of rooms", value=st.session_state.rooms)
    st.session_state.floorHeight = st.text_input("Room/floor to floor height in meter", value=st.session_state.floorHeight)

col_prob1, col_prob2 = st.columns(2)
with col_prob1:
    st.session_state.offerDate = st.text_input("Estimated date to close the offer", value=st.session_state.offerDate, placeholder="e.g. 15.08.2026")
with col_prob2:
    st.session_state.probability = st.text_input("Probability %", value=st.session_state.probability)

st.subheader("Calculation Information & Routing")

col_calc1, col_calc2 = st.columns(2)
with col_calc1:
    st.session_state.calcDrainage = st.radio("Drainage system required", [
        "Single Stack System – HULIOT design drawing, BOQ & Quotation required",
        "Two Stack System (soil / waste / vent) – HULIOT design drawing, BOQ & Quotation required- if only architectural drawing is available",
        "Drainage- Same as per Client/ Consultant drawing -BOQ -Quantity required only."
    ], index=0)

    st.write("**BOQ is compared with:**")
    st.session_state.boqCompare = []
    c_pvc = st.checkbox("PVC")
    c_ci = st.checkbox("Cast Iron")
    c_hdpe = st.checkbox("HDPE")
    c_oth = st.checkbox("CPVC/UPVC/OTHER")
    if c_pvc: st.session_state.boqCompare.append("PVC")
    if c_ci: st.session_state.boqCompare.append("Cast Iron")
    if c_hdpe: st.session_state.boqCompare.append("HDPE")
    if c_oth: st.session_state.boqCompare.append("CPVC/UPVC/OTHER")

with col_calc2:
    st.write("**Drainage Pipe routing:**")
    st.session_state.rteCeiling = st.checkbox("Within false ceiling/ ceiling suspended / underslung system", value=st.session_state.rteCeiling)
    st.session_state.rteToiletSunken = st.checkbox("Toilet Sunken option", value=st.session_state.rteToiletSunken)
    if st.session_state.rteToiletSunken:
        st.session_state.rteToiletMM = st.text_input("Toilet Sunken Depth (mm)", value=st.session_state.rteToiletMM)
    
    st.session_state.rteKitchenSunken = st.checkbox("Kitchen Sunken option", value=st.session_state.rteKitchenSunken)
    if st.session_state.rteKitchenSunken:
        st.session_state.rteKitchenMM = st.text_input("Kitchen Sunken Depth (mm)", value=st.session_state.rteKitchenMM)

    st.session_state.rteUtilitySunken = st.checkbox("Utility Sunken option", value=st.session_state.rteUtilitySunken)
    if st.session_state.rteUtilitySunken:
        st.session_state.rteUtilityMM = st.text_input("Utility Sunken Depth (mm)", value=st.session_state.rteUtilityMM)

    st.write("**For water supply Pipe routing:**")
    st.session_state.wsCeilingDrop = st.checkbox("Within false ceiling & Pipe drop ceiling to wall chasing fixture point consider–Internal Toilets", value=st.session_state.wsCeilingDrop)
    st.session_state.wsWallChase = st.checkbox("In wall chasing piping full - Internal toilets", value=st.session_state.wsWallChase)

col_svc1, col_svc2 = st.columns(2)
with col_svc1:
    st.session_state.serviceFloor = st.selectbox("Service floor available", ["", "Yes", "No"], index=["", "Yes", "No"].index(st.session_state.serviceFloor) if st.session_state.serviceFloor in ["", "Yes", "No"] else 0)
    if st.session_state.serviceFloor == "Yes":
        st.session_state.serviceFloorWhere = st.text_input("Where is the service floor located?", value=st.session_state.serviceFloorWhere)

with col_svc2:
    st.session_state.calcFor = st.selectbox("Calculation Design / BOQ for", ["", "one shaft", "all shafts", "including basement"], index=["", "one shaft", "all shafts", "including basement"].index(st.session_state.calcFor) if st.session_state.calcFor in ["", "one shaft", "all shafts", "including basement"] else 0)

st.session_state.notes = st.text_area("Additional Information / Technical Requirements", value=st.session_state.notes)

st.subheader("Generate & Download Verbatim Document")

active_template_source = None
if uploaded_template is not None:
    active_template_source = uploaded_template
elif template_file_path is not None:
    active_template_source = template_file_path

if active_template_source:
    if st.button("Generate Official Huliot Word File", type="primary", use_container_width=True):
        try:
            doc = Document(active_template_source)

            # Dictionary of simple placeholder text values
            text_replacements = {
                # General fields
                "Date:": f"Date: {st.session_state.date}",
                "Bitrix deal ID:": f"Bitrix deal ID: {st.session_state.projectNo}",
                "Huliot Sales Person & No:": f"Huliot Sales Person & No: {st.session_state.salesPerson} {st.session_state.salesPhone}",
                "Project Name & Address:": f"Project Name & Address: {st.session_state.projectName} {st.session_state.projectAddress}",
                "City:": f"City: {st.session_state.city}",
                "Costumer / Client Name & No:": f"Costumer / Client Name & No: {st.session_state.clientName} {st.session_state.clientPhone}",
                "Architect Name & No:": f"Architect Name & No: {st.session_state.architectName} {st.session_state.architectPhone}",
                "MEP Consultant Name & No:": f"MEP Consultant Name & No: {st.session_state.mepName} {st.session_state.mepPhone}",
                
                # Technical metadata
                "No of Typical floors:": f"No of Typical floors: {st.session_state.floors}",
                "No of shafts:": f"No of shafts: {st.session_state.shafts}",
                "No of rooms:": f"No of rooms: {st.session_state.rooms}",
                "Room/floor to floor height in meter:": f"Room/floor to floor height in meter: {st.session_state.floorHeight}",
                "Estimated date to close the offer.": f"Estimated date to close the offer. {st.session_state.offerDate}",
                "Probability ---": f"Probability --- {st.session_state.probability}%",

                # Sunken depths
                "Toilet Sunken- if Sunken How much in mm:": f"Toilet Sunken- if Sunken How much in mm: {st.session_state.rteToiletMM if st.session_state.rteToiletSunken else ''}",
                "Kitchen Sunken- if Sunken How much in mm:": f"Kitchen Sunken- if Sunken How much in mm: {st.session_state.rteKitchenMM if st.session_state.rteKitchenSunken else ''}",
                "Utility Sunken- if Sunken How much in mm:": f"Utility Sunken- if Sunken How much in mm: {st.session_state.rteUtilityMM if st.session_state.rteUtilitySunken else ''}",
                "if yes where:": f"if yes where: {st.session_state.serviceFloorWhere if st.session_state.serviceFloor == 'Yes' else ''}",

                # Summary inputs
                "Toilet is suspended: -": f"Toilet is suspended: - {'Yes (Ceiling Suspended)' if st.session_state.rteCeiling else 'No'}",
                "Kitchen is suspended: -": f"Kitchen is suspended: - {'Yes (Ceiling Suspended)' if st.session_state.rteCeiling else 'No'}",
            }

            # Helper function to safely replace placeholder text while retaining active run styles (fonts, bold, margins, colors, borders)
            def safe_replace(p, old_text, new_text):
                if old_text in p.text:
                    # 1. First, search inside separate runs (keeps formatting entirely pristine)
                    for run in p.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            return True
                    # 2. If split across multiple runs, merge text safely on paragraph-run levels to prevent style loss
                    full_text = "".join(r.text for r in p.runs)
                    if old_text in full_text:
                        new_full_text = full_text.replace(old_text, new_text)
                        if len(p.runs) > 0:
                            p.runs[0].text = new_full_text
                            for r in p.runs[1:]:
                                r.text = ""
                            return True
                return False

            # Helper function to process run text replacements across paragraphs
            def run_replacements(p):
                for old_patt, new_text in text_replacements.items():
                    if old_patt in p.text:
                        safe_replace(p, old_patt, new_text)

            # Apply replacements on main layout paragraphs
            for paragraph in doc.paragraphs:
                run_replacements(paragraph)

            # Map checkbox replacements by parsing specific runs inside paragraphs
            checkbox_map = {
                # Drawings
                "All or Typical Floor drawing": st.session_state.drwTypical,
                "Typical Toilet and kitchen section details": st.session_state.drwToiletKitchen,
                "Podium/ Diversion/ Basement drawing": st.session_state.drwPodium,
                "Full Building section view attached": st.session_state.drwBuilding,
                "Only Architectural layout attached": st.session_state.drwArchOnly,
                "Any Other drawing attached": st.session_state.drwOther,

                # Drainage Pipes
                "Internal Toilets -Ultra Silent": st.session_state.drnIntUS,
                "Internal Toilets -HT Pro": st.session_state.drnIntHT,
                "Vertical/ External stack -Ultra Silent": st.session_state.drnVertUS,
                "Vertical/ External stack -HT Pro": st.session_state.drnVertHT,
                "Podium/Diversion/ Basement – Ultra Silent": st.session_state.drnPodUS,
                "Podium/Diversion/ Basement - HT Pro": st.session_state.drnPodHT,

                # Water supply Pipes
                "Internal Pipe  (PERT/AL/PERT)": st.session_state.wsIntPipePERT,
                "Internal Pipe  (PPR)": st.session_state.wsIntPipePPR,
                "Internal Fittings   PPSU": st.session_state.wsIntFitPPSU,
                "Internal Fittings   BRASS": st.session_state.wsIntFitBrass,
                "External Pipe  (PERT/AL/PERT)": st.session_state.wsExtPipePERT,
                "External Pipe  (PPR)": st.session_state.wsExtPipePPR,
                "External Fittings   PPSU": st.session_state.wsExtFitPPSU,
                "External Fittings   BRASS": st.session_state.wsExtFitBrass,
                "Terrace looping Pipe  (PERT/AL/PERT)": st.session_state.wsTerPipePERT,
                "Terrace looping Pipe  (PPR)": st.session_state.wsTerPipePPR,
                "Terrace looping Fittings   PPSU": st.session_state.wsTerFitPPSU,
                "Terrace looping Fittings   BRASS": st.session_state.wsTerFitBrass,

                # Priority level
                "A: High": st.session_state.priority == "A: High",
                "B: Medium": st.session_state.priority == "B: Medium",
                "C: Low": st.session_state.priority == "C: Low",

                # Kind of Building
                "Office / commercial building": st.session_state.buildingType == "Office / commercial building",
                "Shopping centre": st.session_state.buildingType == "Shopping centre",
                "Hotel / Restaurant": st.session_state.buildingType == "Hotel / Restaurant",
                "Industrial building": st.session_state.buildingType == "Industrial building",
                "Hospital / nursing home": st.session_state.buildingType == "Hospital / nursing home",
                "School building": st.session_state.buildingType == "School building",
                "Social housing": st.session_state.buildingType == "Social housing",
                "Sport facility": st.session_state.buildingType == "Sport facility",
                "Infrastructure buildings": st.session_state.buildingType == "Infrastructure buildings",
                "Residential building": st.session_state.buildingType == "Residential building",

                # Calculation options
                "Single Stack System": "single-stack" in st.session_state.calcDrainage or "Single Stack" in st.session_state.calcDrainage,
                "Two Stack System": "Two Stack System" in st.session_state.calcDrainage or "Two Stack" in st.session_state.calcDrainage,
                "Drainage- Same as per Client": "Consultant drawing" in st.session_state.calcDrainage or "Same as per" in st.session_state.calcDrainage,

                # Comparison parameters
                "PVC": "PVC" in st.session_state.boqCompare,
                "Cast Iron": "Cast Iron" in st.session_state.boqCompare,
                "HDPE": "HDPE" in st.session_state.boqCompare,
                "CPVC/UPVC/OTHER": "CPVC/UPVC/OTHER" in st.session_state.boqCompare,

                # Routing options
                "Within false ceiling/ ceiling suspended": st.session_state.rteCeiling,
                "Toilet Sunken-": st.session_state.rteToiletSunken,
                "Kitchen Sunken-": st.session_state.rteKitchenSunken,
                "Utility Sunken-": st.session_state.rteUtilitySunken,
                "Within false ceiling & Pipe drop": st.session_state.wsCeilingDrop,
                "In wall chasing piping full": st.session_state.wsWallChase,

                # Calculation target scale
                "one shaft": st.session_state.calcFor == "one shaft",
                "all shafts": st.session_state.calcFor == "all shafts",
                "including basement": st.session_state.calcFor == "including basement",
            }

            # Loop through tables and update paragraphs inside document tables safely
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            run_replacements(paragraph)

            # High fidelity checkbox alignment parser that preserves style
            def process_paragraph_checkboxes(p):
                text_content = p.text
                for marker, is_checked in checkbox_map.items():
                    if marker in text_content:
                        # Find occurrences of empty/filled checkbox symbols and toggle them based on is_checked state
                        if "☐" in text_content or "[ ]" in text_content or "checkbox" in text_content:
                            new_symbol = "☑" if is_checked else "☐"
                            for run in p.runs:
                                if "☐" in run.text:
                                    run.text = run.text.replace("☐", new_symbol)
                                elif "[ ]" in run.text:
                                    run.text = run.text.replace("[ ]", f"[{new_symbol}]")

            # Check inside standard body list structure
            for paragraph in doc.paragraphs:
                process_paragraph_checkboxes(paragraph)

            # Check inside tables cells structures
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph_checkboxes(paragraph)

            # Handle Service Floor Yes/No checkboxes specifically inside body text and tables
            def process_service_floor(p):
                if "Service floor available:" in p.text:
                    for run in p.runs:
                        if "Yes" in run.text and st.session_state.serviceFloor == "Yes":
                            run.text = run.text.replace("☐", "☑")
                        elif "No" in run.text and st.session_state.serviceFloor == "No":
                            run.text = run.text.replace("☐", "☑")

            for paragraph in doc.paragraphs:
                process_service_floor(paragraph)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_service_floor(paragraph)

            # Handle summary text block replacement at the end of the document safely
            for paragraph in doc.paragraphs:
                if "Kitchen / Utility is sunken" in paragraph.text:
                    sunk_summary = []
                    if st.session_state.rteToiletSunken:
                        sunk_summary.append(f"Toilet: {st.session_state.rteToiletMM}mm")
                    if st.session_state.rteKitchenSunken:
                        sunk_summary.append(f"Kitchen: {st.session_state.rteKitchenMM}mm")
                    if st.session_state.rteUtilitySunken:
                        sunk_summary.append(f"Utility: {st.session_state.rteUtilityMM}mm")
                    
                    p_text = "Kitchen / Utility is sunken (How much in mm): - " + (", ".join(sunk_summary) if sunk_summary else "No")
                    safe_replace(paragraph, paragraph.text, p_text)

            # Append the notes paragraph to the doc structure cleanly
            if st.session_state.notes:
                p = doc.add_paragraph()
                p.add_run("\nSpecial Technical Specifications:\n").bold = True
                p.add_run(st.session_state.notes)

            # Output doc representation as binary streams
            bin_stream = io.BytesIO()
            doc.save(bin_stream)
            bin_stream.seek(0)

            # Expose the download button cleanly
            st.download_button(
                label="📥 Download Ready Checklist Document (.docx)",
                data=bin_stream,
                file_name=f"Checklist_Huliot_{(st.session_state.projectName or 'Draft').replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.success("File compilation complete! Click the button above to download.")
        except Exception as err:
            st.error(f"Failed to process file template correctly. Error: {err}")
else:
    st.info("Upload a checklist template (.docx) in the sidebar or commit it to GitHub to compile.")
